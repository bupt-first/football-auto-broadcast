from pathlib import Path

from PIL import Image, ImageOps
from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor
from reportlab.lib import colors
from reportlab.lib.enums import TA_CENTER, TA_LEFT
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import ParagraphStyle, getSampleStyleSheet
from reportlab.lib.units import cm
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.platypus import Image as PdfImage
from reportlab.platypus import PageBreak, Paragraph, SimpleDocTemplate, Spacer, Table, TableStyle


ROOT = Path(__file__).resolve().parent
CROPS = ROOT / "crops"
CROPS.mkdir(parents=True, exist_ok=True)

SOURCE_IMAGES = {
    "circuit": Path(r"E:\QQ\IMG_20260601_151507_edit_87197.jpg"),
    "diagram": Path(r"E:\QQ\Cache_4fe6a32ce0a4d28c.png"),
    "panel": Path(r"E:\QQ\Cache_-7872fee2e693616b.png"),
    "red_curve": Path(r"E:\QQ\Cache_224be428868852f.jpg"),
    "blue_curve": Path(r"E:\QQ\Cache_e6069c46d971460.jpg"),
    "red_table": Path(r"E:\co_wechat\xwechat_files\wxid_vdvrne1fvcyk22_9456\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\91b235be5d3dc3efc54d161cdfe149c2.jpg"),
    "blue_table": Path(r"E:\co_wechat\xwechat_files\wxid_vdvrne1fvcyk22_9456\temp\RWTemp\2026-06\9e20f478899dc29eb19741386f9343c8\295025e24b4c6138c47f4c72573e6ea8.jpg"),
}


FIGURES = [
    {
        "key": "circuit",
        "title": "电路连接",
        "crop": (260, 70, 3790, 2240),
        "width_cm": 16.4,
        "note": "图中展示了 NI myDAQ 接线盒、面包板、限流电阻、发光二极管及连接导线。实验使用采集卡电源输出作为工作电源，并通过定值电阻两端电压换算 LED 电流。",
    },
    {
        "key": "diagram",
        "title": "程序框图",
        "crop": (255, 265, 1165, 975),
        "width_cm": 15.2,
        "note": "程序框图包含 DAQ 助手数据采集、信号均值测量、电压/电流数值换算与 XY 图生成模块，用于自动记录 LED 正向伏安特性曲线。",
    },
    {
        "key": "panel",
        "title": "前面板",
        "crop": (70, 105, 1010, 590),
        "width_cm": 15.5,
        "note": "前面板设置了电压表、电流表、XY 图显示区，以及测量按钮和电源开关，可实时观察采集电压、电流和伏安曲线变化。",
    },
    {
        "key": "red_curve",
        "title": "红光的伏安特性曲线",
        "crop": (70, 170, 1030, 695),
        "width_cm": 15.5,
        "note": "红光 LED 的伏安曲线在较低正向电压附近开始明显上升，随后电流随电压增加快速增大，体现二极管正向导通特性。",
    },
    {
        "key": "blue_curve",
        "title": "蓝光的伏安特性曲线",
        "crop": (55, 160, 1020, 630),
        "width_cm": 15.5,
        "note": "蓝光 LED 的伏安曲线在较高正向电压附近出现明显上升，说明其导通电压相对红光 LED 更高。",
    },
    {
        "key": "red_table",
        "title": "红光的 U-I 记录表",
        "crop": (0, 2120, 2220, 2820),
        "width_cm": 16.3,
        "note": "记录表列出了红光 LED 在不同电压下的电流测量值，为绘制红光正向伏安特性曲线提供原始数据依据。",
    },
    {
        "key": "blue_table",
        "title": "蓝光的 U-I 记录表",
        "crop": (0, 70, 1040, 390),
        "width_cm": 16.3,
        "note": "记录表列出了蓝光 LED 在不同电压下的电流测量值，可与蓝光曲线截图对应检查数据趋势。",
    },
]


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_border(cell, color="D9E2F3"):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), "6")
        element.set(qn("w:color"), color)


def set_table_width(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, width in enumerate(widths_cm):
            row.cells[idx].width = Cm(width)


def set_paragraph_spacing(paragraph, before=0, after=6, line=1.1):
    paragraph.paragraph_format.space_before = Pt(before)
    paragraph.paragraph_format.space_after = Pt(after)
    paragraph.paragraph_format.line_spacing = line


def style_run(run, font="宋体", size=11, bold=False, color=None):
    run.font.name = font
    run._element.rPr.rFonts.set(qn("w:eastAsia"), font)
    run.font.size = Pt(size)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_text(paragraph, text, font="宋体", size=11, bold=False, color=None):
    run = paragraph.add_run(text)
    style_run(run, font=font, size=size, bold=bold, color=color)
    return run


def crop_image(fig):
    source = SOURCE_IMAGES[fig["key"]]
    image = Image.open(source).convert("RGB")
    cropped = image.crop(fig["crop"])
    cropped = ImageOps.exif_transpose(cropped)
    output = CROPS / f"{fig['key']}.jpg"
    cropped.save(output, quality=94, optimize=True)
    return output


def make_contact_sheet(crop_paths):
    thumbs = []
    for fig, path in crop_paths:
        im = Image.open(path).convert("RGB")
        im.thumbnail((520, 320))
        canvas = Image.new("RGB", (560, 380), "white")
        canvas.paste(im, ((560 - im.width) // 2, 12))
        thumbs.append(canvas)
    sheet = Image.new("RGB", (1120, 1520), "white")
    for idx, thumb in enumerate(thumbs):
        x = (idx % 2) * 560
        y = (idx // 2) * 380
        sheet.paste(thumb, (x, y))
    sheet_path = ROOT / "crop_contact_sheet.jpg"
    sheet.save(sheet_path, quality=92)
    return sheet_path


def setup_document(doc):
    section = doc.sections[0]
    section.page_width = Cm(21)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.3)
    section.bottom_margin = Cm(2.2)
    section.left_margin = Cm(2.35)
    section.right_margin = Cm(2.35)
    section.header_distance = Cm(1.2)
    section.footer_distance = Cm(1.2)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "宋体"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
    normal.font.size = Pt(11)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.1

    for name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[name]
        style.font.name = "宋体"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    add_text(footer, "伏安特性曲线的自动测量", size=9, color="666666")


def add_cover(doc):
    for _ in range(3):
        p = doc.add_paragraph()
        set_paragraph_spacing(p, after=0)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(title, after=8)
    add_text(title, "大学物理实验", size=22, bold=True, color="0B2545")

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(subtitle, after=24)
    add_text(subtitle, "伏安特性曲线的自动测量", size=24, bold=True, color="2E74B5")

    info = doc.add_table(rows=5, cols=2)
    info.alignment = WD_TABLE_ALIGNMENT.CENTER
    set_table_width(info, [4.0, 8.6])
    rows = [
        ("姓名", "罗天富"),
        ("学号", "2025211910"),
        ("班级", "202521130"),
        ("实验时间", "2026 年 6 月 1 日"),
        ("实验组号", "16"),
    ]
    for idx, (label, value) in enumerate(rows):
        left, right = info.rows[idx].cells
        for cell in (left, right):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_border(cell)
            for p in cell.paragraphs:
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                set_paragraph_spacing(p, after=0)
        set_cell_shading(left, "E8EEF5")
        add_text(left.paragraphs[0], label, size=12, bold=True, color="1F4D78")
        add_text(right.paragraphs[0], value, size=12)

    note = doc.add_paragraph()
    note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(note, before=20, after=4)
    add_text(note, "截图汇总与数据记录整理", size=12, color="555555")

    doc.add_page_break()


def add_section_heading(doc, title):
    p = doc.add_paragraph(style="Heading 1")
    add_text(p, title, size=16, bold=True, color="2E74B5")


def add_figure(doc, number, fig, image_path):
    note = doc.add_paragraph()
    set_paragraph_spacing(note, before=2, after=4, line=1.15)
    add_text(note, "说明：", bold=True, color="1F4D78")
    add_text(note, fig["note"])

    pic = doc.add_paragraph()
    pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(pic, after=2)
    pic.add_run().add_picture(str(image_path), width=Cm(fig["width_cm"]))

    caption = doc.add_paragraph()
    caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
    set_paragraph_spacing(caption, after=10)
    add_text(caption, f"图 {number}  {fig['title']}", size=10, bold=True, color="333333")


def build_document(crop_paths):
    doc = Document()
    setup_document(doc)
    add_cover(doc)

    add_section_heading(doc, "一、实验截图与程序界面")
    for idx in range(0, 3):
        add_figure(doc, idx + 1, FIGURES[idx], crop_paths[idx][1])

    doc.add_page_break()
    add_section_heading(doc, "二、伏安特性曲线")
    for idx in range(3, 5):
        add_figure(doc, idx + 1, FIGURES[idx], crop_paths[idx][1])

    doc.add_page_break()
    add_section_heading(doc, "三、U-I 原始记录表")
    for idx in range(5, 7):
        add_figure(doc, idx + 1, FIGURES[idx], crop_paths[idx][1])

    output = ROOT / "伏安特性曲线的自动测量_截图汇总.docx"
    doc.save(output)
    return output


def register_pdf_font():
    candidates = [
        Path(r"C:\Windows\Fonts\simsun.ttc"),
        Path(r"C:\Windows\Fonts\simhei.ttf"),
        Path(r"C:\Windows\Fonts\msyh.ttc"),
    ]
    for font_path in candidates:
        if font_path.exists():
            pdfmetrics.registerFont(TTFont("CNFont", str(font_path)))
            return "CNFont"
    return "Helvetica"


def fit_image_size(path, max_width, max_height):
    image = Image.open(path)
    width, height = image.size
    scale = min(max_width / width, max_height / height)
    return width * scale, height * scale


def build_pdf(crop_paths):
    font_name = register_pdf_font()
    pdf_path = ROOT / "伏安特性曲线的自动测量_截图汇总.pdf"
    doc = SimpleDocTemplate(
        str(pdf_path),
        pagesize=A4,
        rightMargin=2.35 * cm,
        leftMargin=2.35 * cm,
        topMargin=2.3 * cm,
        bottomMargin=2.2 * cm,
        title="伏安特性曲线的自动测量_截图汇总",
    )

    styles = getSampleStyleSheet()
    title_style = ParagraphStyle(
        "CoverTitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=24,
        leading=31,
        textColor=colors.HexColor("#2E74B5"),
        alignment=TA_CENTER,
        spaceAfter=18,
    )
    pretitle_style = ParagraphStyle(
        "CoverPretitle",
        parent=styles["Title"],
        fontName=font_name,
        fontSize=22,
        leading=28,
        textColor=colors.HexColor("#0B2545"),
        alignment=TA_CENTER,
        spaceAfter=8,
    )
    h1_style = ParagraphStyle(
        "H1",
        parent=styles["Heading1"],
        fontName=font_name,
        fontSize=16,
        leading=22,
        textColor=colors.HexColor("#2E74B5"),
        spaceBefore=12,
        spaceAfter=8,
    )
    body_style = ParagraphStyle(
        "BodyCN",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=10.5,
        leading=16,
        alignment=TA_LEFT,
        spaceAfter=5,
    )
    caption_style = ParagraphStyle(
        "CaptionCN",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=9.5,
        leading=13,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#333333"),
        spaceAfter=10,
    )
    muted_style = ParagraphStyle(
        "Muted",
        parent=styles["BodyText"],
        fontName=font_name,
        fontSize=11,
        leading=15,
        alignment=TA_CENTER,
        textColor=colors.HexColor("#555555"),
        spaceBefore=18,
    )

    story = [Spacer(1, 2.2 * cm)]
    story.append(Paragraph("大学物理实验", pretitle_style))
    story.append(Paragraph("伏安特性曲线的自动测量", title_style))

    data = [
        ["姓名", "罗天富"],
        ["学号", "2025211910"],
        ["班级", "202521130"],
        ["实验时间", "2026 年 6 月 1 日"],
        ["实验组号", "16"],
    ]
    info_table = Table(data, colWidths=[4.0 * cm, 8.6 * cm], hAlign="CENTER")
    info_table.setStyle(
        TableStyle(
            [
                ("FONTNAME", (0, 0), (-1, -1), font_name),
                ("FONTSIZE", (0, 0), (-1, -1), 12),
                ("ALIGN", (0, 0), (-1, -1), "CENTER"),
                ("VALIGN", (0, 0), (-1, -1), "MIDDLE"),
                ("BACKGROUND", (0, 0), (0, -1), colors.HexColor("#E8EEF5")),
                ("TEXTCOLOR", (0, 0), (0, -1), colors.HexColor("#1F4D78")),
                ("GRID", (0, 0), (-1, -1), 0.6, colors.HexColor("#D9E2F3")),
                ("ROWBACKGROUNDS", (1, 0), (1, -1), [colors.white]),
                ("TOPPADDING", (0, 0), (-1, -1), 8),
                ("BOTTOMPADDING", (0, 0), (-1, -1), 8),
            ]
        )
    )
    story.extend([info_table, Paragraph("截图汇总与数据记录整理", muted_style), PageBreak()])

    sections = [
        ("一、实验截图与程序界面", range(0, 3)),
        ("二、伏安特性曲线", range(3, 5)),
        ("三、U-I 原始记录表", range(5, 7)),
    ]
    page_width = A4[0] - 4.7 * cm
    for section_idx, (heading, indexes) in enumerate(sections):
        if section_idx:
            story.append(PageBreak())
        story.append(Paragraph(heading, h1_style))
        for idx in indexes:
            fig = FIGURES[idx]
            image_path = crop_paths[idx][1]
            story.append(Paragraph(f"<b>说明：</b>{fig['note']}", body_style))
            max_height = 11.2 * cm if idx < 5 else 8.2 * cm
            width, height = fit_image_size(image_path, page_width, max_height)
            story.append(PdfImage(str(image_path), width=width, height=height, hAlign="CENTER"))
            story.append(Paragraph(f"图 {idx + 1}  {fig['title']}", caption_style))

    def footer(canvas, _doc):
        canvas.saveState()
        canvas.setFont(font_name, 8.5)
        canvas.setFillColor(colors.HexColor("#666666"))
        canvas.drawCentredString(A4[0] / 2, 1.2 * cm, "伏安特性曲线的自动测量")
        canvas.restoreState()

    doc.build(story, onFirstPage=footer, onLaterPages=footer)
    return pdf_path


def main():
    crop_paths = []
    for fig in FIGURES:
        crop_paths.append((fig, crop_image(fig)))
    contact_sheet = make_contact_sheet(crop_paths)
    docx_path = build_document(crop_paths)
    pdf_path = build_pdf(crop_paths)
    print(f"DOCX={docx_path}")
    print(f"PDF={pdf_path}")
    print(f"CONTACT={contact_sheet}")


if __name__ == "__main__":
    main()
