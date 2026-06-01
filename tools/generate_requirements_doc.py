from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Inches, Pt, RGBColor


OUT = "doc/football_auto_broadcast_需求分析.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_width(cell, width_dxa):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:type"), "dxa")
    tc_w.set(qn("w:w"), str(width_dxa))


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.find(qn("w:tcMar"))
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def format_table(table, widths=None, header_fill="F2F4F7"):
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    for r_idx, row in enumerate(table.rows):
        for c_idx, cell in enumerate(row.cells):
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)
            if widths:
                set_cell_width(cell, widths[c_idx])
                cell.width = Inches(widths[c_idx] / 1440)
            for p in cell.paragraphs:
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.name = "Microsoft YaHei"
                    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
                    run.font.size = Pt(9)
            if r_idx == 0:
                set_cell_shading(cell, header_fill)
                for p in cell.paragraphs:
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in p.runs:
                        run.bold = True


def set_table_width(table, width_dxa=9360, indent_dxa=120):
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.find(qn("w:tblW"))
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(width_dxa))
    tbl_ind = tbl_pr.find(qn("w:tblInd"))
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:type"), "dxa")
    tbl_ind.set(qn("w:w"), str(indent_dxa))


def add_table(doc, headers, rows, widths):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    set_table_width(table)
    for idx, header in enumerate(headers):
        table.rows[0].cells[idx].text = header
    for row in rows:
        cells = table.add_row().cells
        for idx, value in enumerate(row):
            cells[idx].text = value
    format_table(table, widths)
    doc.add_paragraph()
    return table


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        p.add_run(item)


def add_numbered(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Number")
        p.add_run(item)


def add_kv_table(doc, rows):
    return add_table(
        doc,
        ["项目", "内容"],
        rows,
        [1900, 7460],
    )


def add_callout(doc, title, body):
    table = doc.add_table(rows=1, cols=1)
    table.style = "Table Grid"
    set_table_width(table)
    cell = table.cell(0, 0)
    set_cell_shading(cell, "F4F6F9")
    set_cell_margins(cell, top=120, bottom=120, start=160, end=160)
    p = cell.paragraphs[0]
    run = p.add_run(title)
    run.bold = True
    run.font.color.rgb = RGBColor(31, 77, 120)
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    p.add_run("\n" + body)
    for paragraph in cell.paragraphs:
        for run in paragraph.runs:
            run.font.name = "Microsoft YaHei"
            run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
            run.font.size = Pt(10.5)
    doc.add_paragraph()


def style_document(doc):
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    normal.font.size = Pt(10.5)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.10

    for style_name, size, color, before, after in [
        ("Heading 1", 16, "2E74B5", 16, 8),
        ("Heading 2", 13, "2E74B5", 12, 6),
        ("Heading 3", 12, "1F4D78", 8, 4),
    ]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)

    for style_name in ["List Bullet", "List Number"]:
        style = styles[style_name]
        style.font.name = "Microsoft YaHei"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
        style.font.size = Pt(10.5)
        style.paragraph_format.space_after = Pt(6)
        style.paragraph_format.line_spacing = 1.167


def add_footer(doc):
    section = doc.sections[0]
    footer = section.footer.paragraphs[0]
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("football_auto_broadcast 需求分析")
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(85, 85, 85)


def build_doc():
    doc = Document()
    style_document(doc)
    add_footer(doc)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("football_auto_broadcast 项目需求分析")
    run.bold = True
    run.font.name = "Microsoft YaHei"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    run.font.size = Pt(24)
    run.font.color.rgb = RGBColor(11, 37, 69)

    subtitle = doc.add_paragraph()
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = subtitle.add_run("面向校园、业余与训练赛的低成本、可解释足球自动导播与集锦系统")
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    r.font.size = Pt(12)
    r.font.color.rgb = RGBColor(85, 85, 85)

    add_kv_table(
        doc,
        [
            ("项目名称", "football_auto_broadcast 足球自动转播与集锦生成系统"),
            ("实现语言", "C++17，结合 OpenCV、Qt Widgets、FFmpeg/VideoWriter 输出链路"),
            ("当前形态", "可运行桌面原型：双摄像头导播控制台、目标检测占位算法、手动/自动集锦标记、视频与 JSON 报告输出"),
            ("适用场景", "校园足球、业余比赛、队内训练、低成本固定机位赛事记录"),
            ("核心创新", "不是单纯跟踪画面的 AI 摄像机，而是基于比赛上下文的低成本、可解释自动导演"),
        ],
    )

    add_callout(
        doc,
        "需求定位",
        "本项目面向没有专业转播团队和昂贵场馆基础设施的足球比赛，构建一套能够自动切换视角、生成直播输出、复用事件时间线制作全场集锦和个人集锦，并用可解释指标评价结果质量的 C++ 软件系统。",
    )

    doc.add_heading("1. 项目背景与问题定义", level=1)
    doc.add_paragraph(
        "传统足球转播通常依赖多名摄像师、导播、慢动作回放与字幕人员，画面质量和镜头语言丰富，但成本高、部署复杂，不适合校园和业余比赛。近年的 AI 运动相机能够自动录制、跟踪足球并生成简单片段，但多数产品重点解决“把动作拍进去”，并没有充分解决“理解比赛后如何导播、回放和剪辑”的问题。"
    )
    doc.add_paragraph(
        "football_auto_broadcast 项目以 C++17/OpenCV/Qt 原型为基础，目标是在低成本硬件条件下完成自动导播、集锦生成、个人片段检索和结果评价。系统将实时检测结果保存为事件时间线，直播阶段用于自动切换全景、跟随和特写，赛后阶段复用同一批结构化数据生成全场集锦、球员个人集锦与评价报告。"
    )

    doc.add_heading("2. 现有方案对比", level=1)
    add_table(
        doc,
        ["方案类型", "代表能力", "优势", "不足", "对本项目的启发"],
        [
            (
                "传统专业转播",
                "多机位拍摄、人工导播、慢动作回放、字幕包装",
                "视觉质量高，镜头语言成熟，关键事件捕捉稳定",
                "人员和设备成本高，场馆部署复杂，难以用于校园和训练赛",
                "保留“全景交代战术 + 特写突出事件 + 回放强化记忆”的导播逻辑，但用低成本自动化替代人工流程",
            ),
            (
                "固定机位录制",
                "一台相机记录全场或半场",
                "成本低、部署快、可靠性高",
                "缺少镜头变化，观众体验弱，个人与关键事件检索困难",
                "作为系统的保底输入和战术全景视角，确保任何时刻都有可用画面",
            ),
            (
                "AI 运动相机",
                "自动跟踪球/人、云端直播、简单集锦",
                "安装简单，适合业余赛事，能显著降低人力",
                "通常偏重跟随动作，对攻防态势、禁区威胁、定位球和剪辑解释不足",
                "在“跟踪”基础上增加比赛语义判断和可解释高光评分",
            ),
            (
                "VAR/门线/半自动越位",
                "结构化追踪、位置判断、辅助裁判决策",
                "精度高，数据结构清晰，适合复核关键事实",
                "目标是裁判准确性，不直接服务观众叙事和剪辑",
                "借鉴结构化跟踪数据思想，将其转向导播决策、回放选择和集锦排序",
            ),
        ],
        [1500, 1900, 1900, 2100, 1960],
    )

    doc.add_heading("3. 总体需求", level=1)
    doc.add_heading("3.1 功能性需求", level=2)
    add_numbered(
        doc,
        [
            "视频输入：支持 USB 摄像头、双摄像头和本地比赛视频文件输入；支持摄像头扫描和手动指定全景/特写机位。",
            "目标检测：在原型阶段使用帧差运动检测识别候选球员/足球目标，后续扩展为 YOLO/OpenCV DNN 识别足球、球员、守门员和球门区域。",
            "自动导播：根据目标数量、运动强度、球/人位置、禁区威胁和事件类型，在全景、跟随、特写之间自动切换，并支持人工强制全景或特写。",
            "直播输出：Qt 控制台显示自动播出画面、全景监看和特写监看，录制 panorama_record.mp4、closeup_record.mp4 和 broadcast_record.mp4。",
            "全场集锦：基于射门、扑救、进球、强进攻等事件生成带前后缓冲的 highlight.mp4 和 highlight_report.json。",
            "个人集锦：按球员 ID、人工标注或后续球衣号码/人脸/颜色跟踪生成 personal_highlight_report.json 和个人片段清单。",
            "评价报告：输出事件数量、平均时长、事件密度、目标入镜率、回放价值等指标，并为后续人工标注评测预留精确率、召回率和延迟指标。",
        ],
    )

    doc.add_heading("3.2 非功能性需求", level=2)
    add_table(
        doc,
        ["类别", "需求说明", "验收方式"],
        [
            ("实时性", "导播决策应在 30fps 输入下保持可交互预览，单次模式切换延迟控制在可感知但不突兀的范围内", "记录处理帧率、导播延迟和 UI 卡顿情况"),
            ("稳定性", "摄像头断帧、检测失败或 ROI 缺失时应回退到上一焦点或画面中心，不中断输出", "遮挡/边缘/低光测试仍能保存视频与报告"),
            ("可解释性", "每个高光片段应包含事件类型、时间范围、得分原因和导播理由", "检查 JSON 报告字段完整性"),
            ("可扩展性", "检测、导播、剪辑和 UI 模块之间通过结构体和接口解耦，便于替换模型和新增事件规则", "代码审查模块边界和头文件接口"),
            ("低成本部署", "基础方案使用 1080p USB 摄像头、三脚架和 Windows 笔记本完成演示", "完成 3-5 分钟实地录制测试"),
        ],
        [1600, 5400, 2360],
    )

    doc.add_heading("4. 用户与场景分析", level=1)
    add_table(
        doc,
        ["用户角色", "主要目标", "核心需求"],
        [
            ("普通观众", "看清关键攻防和进球过程", "画面稳定、球和关键球员尽量入镜、切换不突兀、集锦不漏重要事件"),
            ("教练", "复盘阵型、跑位和攻防转换", "保留全景战术信息，支持按事件筛选，输出可度量证据"),
            ("球员", "获取个人表现片段", "个人触球、射门、防守和高光动作可检索，片段前后文完整"),
            ("学生项目评审", "判断方案完整性和实现程度", "看到竞品比较、架构设计、硬件选择、指标体系和 C++ 原型运行结果"),
        ],
        [1500, 2200, 5660],
    )

    doc.add_heading("5. 系统总体方案", level=1)
    doc.add_paragraph("系统采用“采集层 - 感知层 - 导播决策层 - 输出与剪辑层 - 评价层”的分层结构。当前代码已经形成对应模块：")
    add_table(
        doc,
        ["层次", "项目模块", "当前实现", "后续增强"],
        [
            ("视频采集层", "src/video_stream", "VideoStreamManager、DualVideoStreamManager 支持摄像头/文件输入与双机位读取", "增加 RTMP 推流、设备状态恢复、音频同步"),
            ("感知层", "src/detection、src/face_capture", "帧差运动检测、目标框、射门/扑救/进球占位规则、Haar 人脸检测", "引入足球/球员/球门检测模型、SORT/Kalman 跟踪、球衣号码识别"),
            ("导播层", "src/ui/qt_broadcast_window", "Qt 三画面控制台、AUTO/强制全景/强制特写、BroadcastDecision", "加入更平滑的虚拟镜头、定位球和反击场景规则"),
            ("剪辑层", "src/editor", "缓存帧、导入事件、生成 EDL、高光视频、个人报告、JSON 指标", "基于评分排序、去重、BGM/字幕/片头片尾包装"),
            ("评价层", "src/editor/common", "EvaluationMetrics：数量、时长、密度、目标可见性、回放分", "加入人工标注对比、漏检率、冗余率、观众满意度问卷"),
        ],
        [1400, 2200, 3100, 2660],
    )

    doc.add_heading("6. 自动导播需求设计", level=1)
    doc.add_paragraph(
        "自动导播的目标不是机械追球，而是在保证观看舒适性的前提下，根据比赛状态选择最合适的画面。当前原型使用 NORMAL、FOLLOW、CLOSEUP 三种模式，对应全景、跟随和特写。"
    )
    add_table(
        doc,
        ["导播模式", "触发条件", "画面策略", "保底逻辑"],
        [
            ("NORMAL 全景", "目标少、威胁低、需要展示阵型和攻防站位", "优先使用全景机位或较大裁剪范围", "检测失败时回到全景，保证比赛连续性"),
            ("FOLLOW 跟随", "球或主要运动区域持续移动，出现带球推进或攻防转换", "以 ROI 中心平滑裁剪，保持球和邻近球员入镜", "ROI 缺失时沿用上一帧平滑中心"),
            ("CLOSEUP 特写", "射门、扑救、进球、禁区威胁、强身体对抗或手动标记", "短时间切换到特写机位/局部放大，突出动作与表情", "设置保持时间和冷却时间，避免来回跳切"),
        ],
        [1500, 2900, 3300, 1660],
    )
    doc.add_heading("6.1 比赛感知导播规则", level=2)
    add_bullets(
        doc,
        [
            "球位置：球位于中路、边路、禁区或球门附近时采用不同镜头权重。",
            "攻击方向：连续运动向球门区域推进时提高 FOLLOW 或 CLOSEUP 权重。",
            "球员密度：密集对抗区域需要更大视野，避免过度放大导致战术信息丢失。",
            "禁区威胁：球、球员和球门区域同时接近时触发射门/扑救候选事件。",
            "定位球：角球、任意球、点球等场景优先保留全景构图，动作发生后切入特写回放。",
            "反击：高速长距离推进时扩大裁剪窗口并延迟特写，保证传球线路可见。",
        ],
    )

    doc.add_heading("7. 集锦生成需求设计", level=1)
    doc.add_heading("7.1 统一事件时间线", level=2)
    doc.add_paragraph(
        "系统在直播阶段记录每个候选事件的时间戳、类型、目标框、相关目标和导播理由。赛后剪辑不重新从零分析整场视频，而是复用实时事件时间线，减少处理成本，并保证直播导播和赛后剪辑依据一致。"
    )
    add_table(
        doc,
        ["事件字段", "说明", "来源"],
        [
            ("eventType / HighlightType", "射门、扑救、进球、强进攻、手动标记等", "TargetDetectionManager::detectHighlight 与人工 SPACE 标记"),
            ("startSec / endSec", "片段起止时间，包含前缓冲和后缓冲", "EditorConfig preBufferSec/postBufferSec"),
            ("mainTarget / relatedTargets", "主要目标框及相关球员/足球目标", "检测模块输出"),
            ("playerID / belong", "个人集锦归属，当前为占位字段，后续绑定球员身份", "人工标注、球衣号码、人脸或颜色跟踪"),
            ("highlightScore", "可解释高光得分", "VideoEditorManager::scoreHighlight"),
        ],
        [1800, 5200, 2360],
    )

    doc.add_heading("7.2 全场集锦流程", level=2)
    add_numbered(
        doc,
        [
            "读取整场事件列表，过滤置信度低或持续时间不足的候选事件。",
            "按可解释高光分排序，保留进球、射门、扑救、连续进攻和关键防守。",
            "对时间上过近的片段去重或合并，避免同一进攻被重复剪入。",
            "为每个片段扩展前 3 秒和后 5 秒上下文，保证动作因果完整。",
            "生成 EDL 时间线，输出 highlight.mp4、片头、回放标签、边框提示和 JSON 报告。",
        ],
    )

    doc.add_heading("7.3 个人集锦流程", level=2)
    add_numbered(
        doc,
        [
            "建立球员身份来源：演示阶段可通过人工标注或 belong 字段绑定，增强阶段使用球衣号码、人脸和颜色直方图跟踪。",
            "筛选该球员参与的触球、射门、助攻、防守、扑救或跑位片段。",
            "根据球员参与度、动作完整性和事件价值重新排序，而不是只截取进球。",
            "输出个人片段清单和 personal_highlight_report.json，后续可扩展为独立个人视频。",
        ],
    )

    doc.add_heading("7.4 可解释高光评分", level=2)
    add_table(
        doc,
        ["评分因子", "含义", "示例权重"],
        [
            ("事件类型", "进球、射门、扑救、关键传球、抢断等基础价值", "25%"),
            ("场地区域", "禁区、球门附近、中路推进等位置价值", "15%"),
            ("进攻威胁", "速度、推进方向、与球门距离和防守压力", "20%"),
            ("球员参与", "目标球员是否直接触球或参与关键动作", "15%"),
            ("动作连续性", "片段是否包含动作前因、高潮和结果", "10%"),
            ("比分影响", "进球、扳平、反超或关键扑救的比赛影响", "10%"),
            ("回放价值", "动作观赏性、表情反应、是否值得特写回看", "5%"),
        ],
        [1700, 5900, 1760],
    )

    doc.add_heading("8. 评价指标体系", level=1)
    add_table(
        doc,
        ["指标类别", "指标", "定义/计算方式", "意义"],
        [
            ("直播画面质量", "球入镜率", "足球中心或目标框位于输出画面的帧数 / 总帧数", "衡量是否看得见比赛核心"),
            ("直播画面质量", "关键球员入镜率", "事件相关球员出现在输出画面的帧数 / 事件帧数", "衡量导播是否覆盖关键参与者"),
            ("事件覆盖", "漏检关键事件数", "人工标注关键事件中未被系统捕捉的数量", "衡量比赛理解能力"),
            ("导播体验", "切换平滑度", "单位时间切换次数、相邻视角变化幅度和观众评分综合", "避免频繁跳切和眩晕"),
            ("战术保留", "战术信息保留率", "进攻组织/防守站位可见帧数 / 相关片段帧数", "兼顾教练复盘价值"),
            ("集锦质量", "高光相关性", "入选片段中人工认为有效的片段比例", "衡量剪辑精准度"),
            ("集锦质量", "片段冗余率", "重复表达同一事件或无效等待时长 / 集锦总时长", "控制视频节奏"),
            ("效率", "处理延迟", "检测、决策、渲染、写入链路总耗时", "评估实时可用性"),
            ("主观评价", "观众满意度", "从清晰度、完整性、节奏、回放价值等维度打分", "补充客观指标无法覆盖的观看体验"),
        ],
        [1300, 1600, 4200, 2260],
    )

    doc.add_heading("9. 硬件选择与部署需求", level=1)
    add_table(
        doc,
        ["硬件", "推荐规格", "用途", "选择理由"],
        [
            ("全景摄像头", "1080p/30fps USB 摄像头或广角相机", "记录半场/全场战术视角", "低成本、易部署，是自动导播的保底画面"),
            ("特写摄像头", "1080p/30fps USB 摄像头，较窄视角", "补充射门、庆祝、扑救等局部画面", "当前 Qt 原型已支持双摄输入和特写监看"),
            ("三脚架/夹具", "1.5m-2.2m 稳定支架", "固定机位并减少画面抖动", "提升检测稳定性，降低误报"),
            ("计算设备", "Windows 笔记本，Visual Studio、CMake、OpenCV、Qt 环境", "运行 C++ 程序、录制视频、生成报告", "与当前项目构建和演示环境一致"),
            ("存储与供电", "充足硬盘空间、移动电源或延长线", "长时间比赛录制", "避免中途断电和录像失败"),
        ],
        [1500, 2300, 2400, 3160],
    )
    doc.add_paragraph(
        "推荐部署位置为边线中场附近，摄像头略向下俯拍。小场比赛至少覆盖半场，大场比赛优先保证攻防主要区域可见。测试时应覆盖静态摄像头、手动高光、自动高光、完整报告和个人报告五类用例。"
    )

    doc.add_heading("10. 软件实现需求与当前项目映射", level=1)
    add_table(
        doc,
        ["需求", "当前代码支持", "验收输出"],
        [
            ("C++ 实现", "CMakeLists.txt 构建 C++17 程序，主入口 src/main.cpp", "football_auto_broadcast.exe 可运行"),
            ("双机位导播", "DualVideoStreamManager + QtBroadcastWindow 三画面控制台", "broadcast_record.mp4、panorama_record.mp4、closeup_record.mp4"),
            ("检测与事件", "TargetDetectionManager 输出 TargetInfo 和 HighlightInfo", "检测框、自动高光候选、控制台指标变化"),
            ("手动兜底", "SPACE 标记高光，AUTO/强制全景/强制特写控制", "重要事件不会完全依赖占位检测算法"),
            ("视频剪辑", "VideoEditorManager 缓存帧、导出 highlight.mp4 和报告", "highlight.mp4、highlight_report.json"),
            ("个人报告", "exportPersonal 与 player/belong 字段", "personal_highlight_report.json"),
            ("指标评价", "EvaluationMetrics 和 JSON 字段", "集锦数量、平均时长、事件密度、目标可见性、回放分"),
        ],
        [2100, 5000, 2260],
    )

    doc.add_heading("11. 开发计划与优先级", level=1)
    add_table(
        doc,
        ["阶段", "目标", "主要任务", "成果"],
        [
            ("阶段 1：稳定原型", "保证演示可运行", "完善双摄打开、录制、手动标记、报告输出和异常处理", "可演示的 C++ 桌面程序"),
            ("阶段 2：比赛语义增强", "从运动检测升级到足球理解", "接入球/人/球门检测模型，增加禁区、反击、定位球规则", "更可靠的自动导播决策"),
            ("阶段 3：集锦质量提升", "提高剪辑可用性", "实现评分排序、去重、片段合并、个人身份绑定", "全场与个人高光视频/报告"),
            ("阶段 4：评测与答辩", "形成可量化证据", "人工标注测试片段，对比固定机位/手动导播/自动导播", "指标表、截图、演示视频、PPT 材料"),
        ],
        [1500, 2300, 3600, 1960],
    )

    doc.add_heading("12. 风险与约束", level=1)
    add_table(
        doc,
        ["风险", "影响", "应对策略"],
        [
            ("检测算法仍为占位", "自动事件识别可能漏检或误检", "保留 SPACE 手动标记；答辩中明确原型与后续模型替换计划"),
            ("单机位视角有限", "无法同时兼顾战术全景和表情特写", "使用双摄方案：全景为主、特写补充；无特写时使用虚拟裁剪"),
            ("光照和抖动", "帧差检测产生误报，目标框不稳定", "固定三脚架、避免逆光、加入 ROI 平滑和冷却时间"),
            ("个人身份识别困难", "个人集锦归属不稳定", "先支持人工/球员 ID 标注，再扩展球衣号码和颜色跟踪"),
            ("实时性能", "模型升级后可能降低帧率", "分辨率降采样、异步检测、只在关键帧运行重模型"),
        ],
        [2100, 3200, 4060],
    )

    doc.add_heading("13. 验收标准", level=1)
    add_bullets(
        doc,
        [
            "能够说明并比较传统转播、固定机位、AI 运动相机和裁判辅助技术的优缺点。",
            "能够展示自动导播方案：输入、检测、导播决策、输出、录制和报告链路完整。",
            "能够生成全场集锦和个人集锦相关报告，且解释每个片段的入选原因。",
            "能够给出覆盖画面质量、事件覆盖、导播体验、战术保留、集锦质量和实时性能的评价指标体系。",
            "能够给出低成本硬件部署方案，并完成至少 3-5 分钟足球场景测试。",
            "软件以 C++ 实现，项目可构建、可运行，并能输出视频文件和 JSON 结果。",
        ],
    )

    doc.add_heading("14. 结论", level=1)
    doc.add_paragraph(
        "football_auto_broadcast 的需求核心是构建一个面向低成本足球赛事的自动导演系统：它既要能在比赛中自动选择视角、保持关键目标入镜，又要在赛后复用事件时间线生成全场和个人集锦，并用可解释指标证明输出质量。当前 C++ 项目已经具备视频输入、双摄控制台、基础检测、自动/手动高光、视频导出和 JSON 评价报告等原型能力，后续重点是增强足球语义检测、完善导播规则、提升个人身份绑定和建立人工标注评测集。"
    )

    doc.save(OUT)


if __name__ == "__main__":
    build_doc()
